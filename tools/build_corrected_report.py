from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "LinguaLeap_AI_Implementation_Aligned_Report.docx"
ASSETS = ROOT / "docs" / "report_assets"

NAVY = RGBColor(31, 58, 95)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(35, 42, 52)
MUTED = RGBColor(92, 102, 112)
LIGHT = "F4F6F9"
HEADER_FILL = "E8EEF5"
GREEN = RGBColor(45, 112, 78)
GOLD = RGBColor(122, 90, 0)


def set_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])


def configure_page(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_run = header.add_run("LINGUALEAP AI | IMPLEMENTATION-ALIGNED PROJECT REPORT")
        set_font(header_run, size=8.5, color=MUTED, bold=True)
        add_page_field(section.footer.paragraphs[0])


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_font(run, size=9.5, color=NAVY, bold=True)
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().append(repeat)
    for row_data in rows:
        row = table.add_row()
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for index, value in enumerate(row_data):
            p = row.cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            set_font(run, size=9.5)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(52)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PIMPRI CHINCHWAD UNIVERSITY")
    set_font(r, size=13, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("LINGUALEAP AI")
    set_font(r, size=30, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("Intelligent Language Learning Platform with Moderator-Aware Evaluation")
    set_font(r, size=15, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Implementation-Aligned Project Report")
    set_font(r, size=11, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(52)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Submitted toward the partial fulfillment of the requirements of\nBachelor of Technology in Computer Science and Engineering")
    set_font(r, size=11)

    members = [
        "Jai Bonde | PRN: SOE23201040063",
        "Atharva Mejari | PRN: SOE23201040076",
        "Anuja Bahirat | PRN: SOE23201060004",
        "Atharva Salunke | PRN: SOE23201040024",
    ]
    for member in members:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(member)
        set_font(r, size=10.5, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(32)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Department of Computer Science & Engineering\nAcademic Year 2025-26")
    set_font(r, size=10.5, color=MUTED)
    doc.add_page_break()


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        first, rest = text.split(bold_lead, 1)
        r = p.add_run(first + bold_lead)
        set_font(r, bold=True)
        r = p.add_run(rest)
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r)


def add_numbered(doc, items):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_node])
        p._p.get_or_add_pPr().append(num_pr)
        r = p.add_run(item)
        set_font(r)


def add_callout(doc, label, text, color=GREEN):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "5")
        border.set(qn("w:color"), "A7ADB5")
        borders.append(border)
    p_pr.append(borders)
    r = p.add_run(label.upper() + "\n")
    set_font(r, size=9, color=color, bold=True)
    r = p.add_run(text)
    set_font(r, size=10.5)


def make_diagram(path, title, columns):
    image = Image.new("RGB", (1600, 720), "white")
    draw = ImageDraw.Draw(image)
    font_dir = Path("C:/Windows/Fonts")
    title_font = ImageFont.truetype(str(font_dir / "arialbd.ttf"), 42)
    box_font = ImageFont.truetype(str(font_dir / "arialbd.ttf"), 28)
    item_font = ImageFont.truetype(str(font_dir / "arial.ttf"), 22)
    draw.text((60, 40), title, font=title_font, fill=(31, 58, 95))
    x_positions = [70, 565, 1060]
    colors = [(232, 238, 245), (237, 248, 230), (255, 248, 216)]
    for idx, (heading, items) in enumerate(columns):
        x = x_positions[idx]
        draw.rounded_rectangle((x, 145, x + 430, 620), radius=24, fill=colors[idx], outline=(80, 100, 120), width=3)
        draw.text((x + 28, 178), heading, font=box_font, fill=(31, 58, 95))
        y = 245
        for item in items:
            draw.ellipse((x + 28, y + 8, x + 40, y + 20), fill=(45, 112, 78))
            draw.text((x + 56, y), item, font=item_font, fill=(35, 42, 52))
            y += 62
        if idx < 2:
            draw.line((x + 440, 380, x + 490, 380), fill=(46, 116, 181), width=6)
            draw.polygon([(x + 490, 380), (x + 466, 366), (x + 466, 394)], fill=(46, 116, 181))
    image.save(path)


def build():
    ASSETS.mkdir(parents=True, exist_ok=True)
    architecture = ASSETS / "architecture.png"
    evaluation = ASSETS / "evaluation.png"
    make_diagram(
        architecture,
        "LinguaLeap AI - Implemented System Architecture",
        [
            ("React Frontend", ["Authentication", "Lessons", "Practice tools", "Analytics UI", "Privacy controls"]),
            ("FastAPI Backend", ["JWT authorization", "Learning APIs", "Groq integration", "Metric computation", "CSV / JSON export"]),
            ("Persistence", ["SQLAlchemy models", "Alembic migrations", "SQLite local", "Neon PostgreSQL production", "Consent metadata"]),
        ],
    )
    make_diagram(
        evaluation,
        "Consent-Based Evaluation Design",
        [
            ("Assignment", ["LLM tutor", "Structured baseline", "Text-only", "Multimodal"]),
            ("Moderators", ["Proficiency", "Language / skill", "Task complexity", "Feedback type", "Measured duration"]),
            ("Evidence", ["Observed records", "Pre/post scores", "Spaced reviews", "Simulated labels", "Anonymized export"]),
        ],
    )

    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_title_page(doc)

    add_heading(doc, "Revision Status and Claim Boundary", 1)
    add_callout(
        doc,
        "Purpose of this revision",
        "This report replaces unsupported implementation and effectiveness claims in the earlier 138-page draft. It documents the software that exists as of 12 June 2026 and treats the research component as an evaluation design, not a completed or validated study.",
        GOLD,
    )
    add_table(
        doc,
        ["Area", "Earlier draft implied", "Implementation-aligned statement"],
        [
            ("Product identity", "A meta-analytic research engine", "An intelligent language-learning platform with a secondary moderator-aware research dashboard"),
            ("AI", "Custom deep-learning models trained with TensorFlow/PyTorch", "Hosted Groq language model and Whisper speech-to-text; no custom model training"),
            ("Database", "Hybrid PostgreSQL and MongoDB", "SQLAlchemy relational schema; SQLite locally; managed Neon PostgreSQL in production with Alembic migrations"),
            ("Infrastructure", "Docker/Kubernetes and large-scale cloud processing", "React and FastAPI deployed as Vercel applications"),
            ("Results", "Empirical validation and demonstrated effectiveness", "Functional verification only; no adequately powered learning-effectiveness study has been completed"),
        ],
        [1800, 3300, 4260],
    )

    add_heading(doc, "Abstract", 1)
    add_para(
        doc,
        "LinguaLeap AI is an intelligent and interactive language-learning platform that supports Hindi, German, and Japanese learners through structured lessons, conversational practice, contextual translation, pronunciation practice, quizzes, and spaced review. The product is designed to make practice feel closer to working with a personal tutor than reading a static lesson. The React frontend presents reusable learner flows, while a FastAPI backend handles authentication, curriculum progress, AI requests, analytics, privacy controls, and research exports."
    )
    add_para(
        doc,
        "The implemented AI layer uses Groq-hosted services. Tutor and translation responses are validated with Pydantic schemas, retried when malformed, protected by prompt-boundary instructions, and replaced with learner-safe fallbacks when the provider is unavailable. Pronunciation practice records microphone audio, transcribes it with Whisper, and calculates a transparent text-similarity score against the target phrase. This score is a practice aid and not a phoneme-level acoustic assessment."
    )
    add_para(
        doc,
        "A consent-based evaluation layer records learner proficiency, language, skill, task complexity, modality, feedback strategy, measured interaction duration, experiment group, and optional pre/post-test scores. Observed and simulated records are separated. Simulated data exists only to demonstrate the dashboard and is excluded from primary research comparisons. Therefore, this report presents the platform as an implemented educational software project with a defensible future experiment design; it does not claim validated learning effectiveness."
    )
    add_heading(doc, "Contents", 1)
    contents = [
        "Introduction and problem statement",
        "Objectives and project scope",
        "Implemented architecture and technology stack",
        "Learner-facing modules",
        "AI integration and reliability",
        "Data model and longitudinal memory",
        "Spaced repetition",
        "Moderator-aware experiment design",
        "Analytics methodology and evidence rules",
        "Research ethics and privacy",
        "Testing and verification",
        "Deployment, persistence, and limitations",
        "Results, future work, and conclusion",
        "Appendix A. API surface",
        "Appendix B. Claim audit",
    ]
    add_numbered(doc, contents)
    doc.add_page_break()

    add_heading(doc, "1. Introduction and Problem Statement", 1)
    add_para(
        doc,
        "Traditional language-learning workflows can become repetitive because they separate reading, conversation, pronunciation, translation, and review into disconnected activities. Learners also receive limited immediate feedback and may not know which vocabulary or grammar patterns need revisiting. LinguaLeap AI addresses this product problem through one connected learning environment."
    )
    add_heading(doc, "1.1 Product Context", 2)
    add_para(
        doc,
        "The main learning direction is English to Hindi, with German and Japanese included to demonstrate reusable curriculum and practice components. The application is learner-first: research controls support responsible evaluation but do not replace the central experience of lessons, practice, feedback, and progress."
    )
    add_heading(doc, "1.2 Problem Statement", 2)
    add_para(
        doc,
        "How can a web-based platform combine structured language content, conversational AI, speaking and translation practice, learner memory, and transparent progress analytics while collecting enough consented context to support future educational evaluation?"
    )
    add_heading(doc, "1.3 Stakeholders", 2)
    add_bullets(doc, [
        "Learners who need engaging and repeatable practice.",
        "Educators who need interpretable progress records rather than opaque AI scores.",
        "Project evaluators who need a truthful demonstration of implemented features.",
        "Authorized researchers who need anonymized, consent-gated interaction data.",
    ])

    add_heading(doc, "2. Objectives and Project Scope", 1)
    add_heading(doc, "2.1 Implemented Objectives", 2)
    add_bullets(doc, [
        "Provide reusable React interfaces for authentication, language selection, lessons, practice, analytics, and privacy.",
        "Support AI tutor conversations with corrections, encouragement, vocabulary memory, and grammar memory.",
        "Provide translation with native script, romanization, and concise cultural notes.",
        "Provide microphone-based pronunciation practice using speech-to-text and transparent similarity scoring.",
        "Track lesson completion, XP, quiz performance, skill performance, and spaced-review state.",
        "Record moderator variables and optional pre/post assessments for consented evaluation.",
        "Separate simulated demonstration records from observed learner records.",
        "Provide participant consent withdrawal, personal data export, and account deletion.",
    ])
    add_heading(doc, "2.2 Out of Scope", 2)
    add_bullets(doc, [
        "Training or fine-tuning a proprietary language model.",
        "Claiming phoneme-level pronunciation diagnosis.",
        "Conducting a completed meta-analysis across published studies.",
        "Claiming statistically validated improvement or production-scale reliability.",
        "Using TensorFlow, PyTorch, MongoDB, Kubernetes, or a hybrid database architecture.",
    ])

    add_heading(doc, "3. Implemented Architecture and Technology Stack", 1)
    architecture_shape = doc.add_picture(str(architecture), width=Inches(6.45))
    architecture_shape._inline.docPr.set(
        "descr",
        "Three-column architecture diagram showing the React frontend flowing to the FastAPI backend and then to SQLAlchemy persistence.",
    )
    architecture_shape._inline.docPr.set("title", "LinguaLeap implemented architecture")
    p = doc.add_paragraph("Figure 1. Implemented system architecture.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)
    add_heading(doc, "3.1 Technology Stack", 2)
    add_table(
        doc,
        ["Layer", "Implemented technology", "Responsibility"],
        [
            ("Frontend", "React, TypeScript, Vite, Zustand, Axios", "Learner interface, persisted authentication state, API integration"),
            ("Backend", "FastAPI, Pydantic, SQLAlchemy", "REST endpoints, validation, business rules, persistence"),
            ("Security", "bcrypt, JWT", "Password hashing and authenticated API access"),
            ("AI services", "Groq-hosted LLM and Whisper", "Tutor, translation, and speech transcription"),
            ("Database", "SQLite local; managed Neon PostgreSQL production", "Accounts, progress, interactions, review schedules"),
            ("Migrations", "Alembic", "Versioned schema creation and future upgrades"),
            ("Deployment", "Vercel", "Public frontend and serverless FastAPI backend"),
        ],
        [1500, 3000, 4860],
    )

    add_heading(doc, "4. Learner-Facing Modules", 1)
    modules = [
        ("Authentication and profile", "Registration captures proficiency, learning goal, optional research consent, and secure credentials. JWT state persists across refreshes."),
        ("Language selection", "Learners select Hindi, German, or Japanese and receive language-specific content, scripts, pronunciation hints, and exercises."),
        ("Structured lessons", "Each language contains three levels with objectives, grammar or usage explanations, native examples, romanization, vocabulary, XP, and enforced prerequisites."),
        ("Listening support", "Lesson examples and pronunciation targets use the browser Speech Synthesis API for model playback."),
        ("Tutor chat", "Learners practice vocabulary, speaking, writing, or creativity with an AI tutor that receives a bounded learner-memory profile."),
        ("Translation", "English phrases are translated into the target language with romanization and cultural or politeness guidance."),
        ("Pronunciation", "The browser records audio; Whisper transcribes it; the backend compares transcript and target phrase and returns a practice score and feedback."),
        ("Exercises", "MCQ and fill-in-the-blank activities store answer, expected answer, correctness, skill, complexity, XP, and measured duration."),
        ("Analytics and privacy", "Learners see personal performance and can record assessments, export their data, change consent, or delete the account."),
    ]
    for title, description in modules:
        add_heading(doc, title, 2)
        add_para(doc, description)

    add_heading(doc, "5. AI Integration and Reliability", 1)
    add_heading(doc, "5.1 Hosted Model Boundary", 2)
    add_para(
        doc,
        "LinguaLeap calls Groq-hosted models through an OpenAI-compatible client. The project does not train, fine-tune, or own the underlying model. Prompts instruct the model to treat learner text as practice content rather than system instructions."
    )
    add_heading(doc, "5.2 Structured Validation", 2)
    add_numbered(doc, [
        "Limit and sanitize the learner message.",
        "Request a JSON response with explicit fields.",
        "Validate the response against a Pydantic schema.",
        "Retry once at lower temperature if parsing or provider execution fails.",
        "Return a bounded fallback lesson when the live provider is unavailable.",
    ])
    add_heading(doc, "5.3 Reliability Limitations", 2)
    add_bullets(doc, [
        "Hosted LLM output can still contain linguistic or cultural mistakes.",
        "The current prompt-injection controls reduce common instruction overrides but are not a complete security sandbox.",
        "Pronunciation similarity depends on transcription quality and should not be interpreted as a clinical or phonetic score.",
        "Educator review remains appropriate for high-stakes assessment.",
    ])

    add_heading(doc, "6. Data Model and Longitudinal Memory", 1)
    add_table(
        doc,
        ["Entity", "Stored information", "Purpose"],
        [
            ("User", "Anonymous ID, profile, consent, groups, assessments, XP", "Authentication, personalization, ethical research controls"),
            ("MemoryProfile", "Notes, vocabulary focus, grammar focus", "Bounded context for tutor responses"),
            ("LessonProgress", "Lesson, language, completion, score, date", "Prerequisites and progress tracking"),
            ("Interaction", "Prompt, response, skill, moderators, score, duration", "Personal analytics and consented research evidence"),
            ("ReviewItem", "Term, quality, difficulty, interval, next review", "Spaced repetition scheduling"),
        ],
        [1600, 4300, 3460],
    )
    add_heading(doc, "6.1 Moderator Capture", 2)
    add_para(
        doc,
        "Every learner-generated interaction records a proficiency snapshot, target language, skill, task complexity, modality, feedback type, experiment group, delivery group, correctness or score, and measured engagement duration. Optional pre/post-test values are copied into relevant assessment records."
    )

    add_heading(doc, "7. Spaced Repetition", 1)
    add_para(
        doc,
        "Vocabulary introduced by lessons, tutor updates, and exercises is stored as a review item. The current scheduler implements a simplified SM-2 process. Learners rate recall quality from 0 to 5. Ratings below 3 reset the repetition count and schedule a near-term retry; successful recall increases the interval using the stored easiness factor."
    )
    add_heading(doc, "7.1 Stored Review State", 2)
    add_bullets(doc, [
        "Term and translation",
        "Language",
        "Recall quality",
        "Difficulty or easiness factor",
        "Repetition count",
        "Current interval in days",
        "Last review timestamp",
        "Next review timestamp",
    ])
    add_callout(
        doc,
        "Interpretation",
        "This is an implemented scheduling algorithm, not evidence that retention improved. Retention claims require longitudinal learner data and an appropriate study design.",
        GOLD,
    )

    add_heading(doc, "8. Moderator-Aware Experiment Design", 1)
    evaluation_shape = doc.add_picture(str(evaluation), width=Inches(6.45))
    evaluation_shape._inline.docPr.set(
        "descr",
        "Three-column evaluation diagram showing consent-based group assignment, recorded moderator variables, and observed or simulated evidence.",
    )
    evaluation_shape._inline.docPr.set("title", "LinguaLeap consent-based evaluation design")
    p = doc.add_paragraph("Figure 2. Consent-based evaluation design.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)
    add_heading(doc, "8.1 Assignment", 2)
    add_para(
        doc,
        "Only learners who provide research consent are assigned. A server-keyed HMAC of the normalized email produces stable assignment without exposing the bucket rule to the client."
    )
    add_table(
        doc,
        ["Factor", "Group A", "Group B", "Implemented behavior"],
        [
            ("Tutor strategy", "LLM tutor", "Structured non-LLM baseline", "Tutor endpoint selects either Groq response or deterministic structured practice"),
            ("Delivery", "Text-only", "Multimodal", "Text-only forces text modality; multimodal permits audio interactions"),
        ],
        [1500, 1800, 2100, 3960],
    )
    add_heading(doc, "8.2 Validity Conditions", 2)
    add_bullets(doc, [
        "Use the same pre/post assessment instrument for each participant.",
        "Define recruitment, sample size, duration, and exclusion criteria before analysis.",
        "Report group sample sizes and uncertainty, not only mean percentages.",
        "Exclude simulated records from observed claims.",
        "Avoid causal language until assignment integrity and statistical assumptions are reviewed.",
    ])

    add_heading(doc, "9. Analytics Methodology and Evidence Rules", 1)
    add_heading(doc, "9.1 Learner Analytics", 2)
    add_para(
        doc,
        "The learner dashboard calculates interaction count, mean score, observed correctness, average measured engagement, completed lessons, due reviews, skill-level means, and an observed recall curve from stored exercise and review outcomes. Cross-group comparisons are not shown as personal conclusions."
    )
    add_heading(doc, "9.2 Research Analytics", 2)
    add_para(
        doc,
        "The role-protected research dashboard computes group counts, mean score, correctness rate, mean engagement duration, and available learning gain from stored consented interactions. Primary comparison rows exclude simulated records. A separate simulated section demonstrates the interface and is explicitly labeled as non-evidence."
    )
    add_heading(doc, "9.3 Metrics That Are Not Claimed", 2)
    add_bullets(doc, [
        "No independently validated model accuracy.",
        "No statistical significance or confidence intervals from an adequately powered study.",
        "No meta-analytic effect size across published investigations.",
        "No proof that the LLM tutor outperforms the structured baseline.",
        "No proof that multimodal delivery outperforms text-only delivery.",
    ])

    add_heading(doc, "10. Research Ethics and Privacy", 1)
    add_bullets(doc, [
        "Research participation is optional and separate from basic account creation.",
        "Participants receive an anonymous identifier.",
        "Consent can be withdrawn; prior interactions are removed from experiment groupings.",
        "Learners can export their own profile, interaction, and review data as JSON.",
        "Learners can delete their account and associated learning records.",
        "Research CSV exports use anonymous IDs and require a researcher role.",
        "Researcher registration requires both an administrative email namespace and a private access code.",
        "Collected data and intended purpose are explained in the privacy interface.",
    ])
    add_heading(doc, "10.1 Remaining Ethics Work", 2)
    add_para(
        doc,
        "Before formal human-participant research, the team should obtain institutional approval where required, finalize an information sheet and consent form, define retention periods, document data-controller responsibilities, and establish a process for participant questions and adverse events."
    )

    add_heading(doc, "11. Testing and Verification", 1)
    add_para(
        doc,
        "As of 12 June 2026, the backend test suite contains 12 focused tests and passes completely. The frontend TypeScript compilation and Vite production build also pass."
    )
    add_table(
        doc,
        ["Verification area", "Evidence"],
        [
            ("Authentication", "Registration, login token, and consent assignment tests"),
            ("Authorization", "Research routes reject learners; researcher code rejects unauthorized signup"),
            ("Lessons", "Prerequisite enforcement and duplicate-XP prevention"),
            ("Exercises", "Correct and incorrect scoring plus computed analytics"),
            ("Engagement", "Stored analytics reflect client-measured seconds"),
            ("Research export", "Consent-gated records and anonymous identifiers"),
            ("Spaced repetition", "Successful interval growth and failed-recall reset"),
            ("Database migration", "Empty database upgrades to Alembic revision 20260612_0001"),
            ("Deployment", "Public health, registration, login, and learner dashboard checks"),
        ],
        [2600, 6760],
    )
    add_callout(
        doc,
        "Test scope",
        "These checks establish implemented behavior and regression protection. They do not validate pedagogical effectiveness or language quality across all supported content.",
        GOLD,
    )

    add_heading(doc, "12. Deployment, Persistence, and Limitations", 1)
    add_heading(doc, "12.1 Current Deployment", 2)
    add_bullets(doc, [
        "Frontend: https://lingualeap-ai-eight.vercel.app",
        "Backend API: https://lingualeap-api.vercel.app",
        "Runtime secrets are stored as encrypted Vercel environment variables.",
        "Local environment files and SQLite databases are excluded from deployment uploads.",
    ])
    add_heading(doc, "12.2 PostgreSQL Status", 2)
    add_para(
        doc,
        "The production backend is connected to managed Neon PostgreSQL through Vercel Marketplace. The application normalizes the pooled PostgreSQL URL, uses the psycopg driver, and applies the Alembic schema migration during startup. Persistence was verified by creating a learner account, deploying a new production version, and successfully authenticating the same account afterward."
    )
    add_heading(doc, "12.3 Current Limitations", 2)
    add_bullets(doc, [
        "The free managed database tier remains subject to provider quotas and operational limits.",
        "Limited lesson breadth compared with mature commercial platforms.",
        "Browser text-to-speech voice quality varies by device and installed voices.",
        "Pronunciation scoring compares text transcriptions rather than acoustic phonemes.",
        "Hosted AI availability, latency, and output quality depend on Groq.",
        "Research assignments and metrics are implemented, but no formal participant study is complete.",
    ])

    add_heading(doc, "13. Results, Future Work, and Conclusion", 1)
    add_heading(doc, "13.1 Verified Project Outcomes", 2)
    add_bullets(doc, [
        "A deployed learner-facing application for Hindi, German, and Japanese.",
        "Integrated lesson, conversation, translation, pronunciation, quiz, and review workflows.",
        "Persistent data models, a versioned migration, and managed production PostgreSQL.",
        "Consent-based moderator capture and defensible separation of observed and simulated records.",
        "Automated tests for high-risk behavioral and research-data paths.",
        "Documentation that explicitly distinguishes implemented technology from future research.",
    ])
    add_heading(doc, "13.2 Recommended Future Work", 2)
    add_numbered(doc, [
        "Add scheduled database backups and a documented restoration drill.",
        "Expand each language with additional CEFR-aligned levels, listening assessments, and educator-reviewed content.",
        "Replace transcript similarity with a phoneme-aware pronunciation service if higher-fidelity feedback is required.",
        "Add frontend component and end-to-end tests for microphone permissions, lesson playback, and data-export flows.",
        "Create an approved study protocol, recruit participants, and perform sample-size planning before effectiveness analysis.",
        "Report uncertainty, missing-data rules, and statistical assumptions in any future experimental results.",
    ])
    add_heading(doc, "13.3 Conclusion", 2)
    add_para(
        doc,
        "LinguaLeap AI now matches the project story presented in interviews: it is an interactive language-learning platform built with React, API integration, reusable learner flows, AI-assisted practice, and progress tracking. Its research layer is a carefully bounded extension that records moderators and supports future comparisons. The implemented system is substantial, demonstrable, and testable without relying on claims about custom deep learning, unsupported infrastructure, or unvalidated educational outcomes."
    )

    add_heading(doc, "Appendix A. Implemented API Surface", 1)
    add_table(
        doc,
        ["Area", "Representative endpoints"],
        [
            ("Authentication", "POST /auth/register; POST /auth/login; GET /auth/me"),
            ("Learner setup", "POST /platform/language; GET /platform/bootstrap"),
            ("Lessons", "GET /platform/curriculum; GET /platform/lessons/{id}; POST /platform/lessons/{id}/complete"),
            ("Practice", "GET /platform/exercises; POST /platform/exercises/submit; POST /platform/tutor"),
            ("Language tools", "POST /platform/translate; POST /platform/pronunciation"),
            ("Spaced review", "GET /platform/reviews; POST /platform/reviews/{id}"),
            ("Assessment", "POST /platform/assessment"),
            ("Analytics", "GET /platform/analytics; GET /platform/research"),
            ("Ethics", "GET /platform/privacy; PATCH /platform/privacy/consent; GET /platform/account/export; DELETE /platform/account"),
            ("Research export", "GET /platform/research/export"),
        ],
        [2100, 7260],
    )

    add_heading(doc, "Appendix B. Claim Audit", 1)
    add_table(
        doc,
        ["Claim in earlier report", "Status", "Correct replacement"],
        [
            ("TensorFlow/PyTorch models are implemented and trained", "Remove", "Groq-hosted model is called through an API; no custom training"),
            ("Hugging Face, spaCy, and NLTK power the analysis", "Remove", "Not present in the implemented dependency set"),
            ("Hybrid MongoDB/PostgreSQL architecture", "Remove", "SQLAlchemy relational schema; SQLite local; managed Neon PostgreSQL production"),
            ("Docker and Kubernetes provide scalable deployment", "Remove", "Vercel hosts the React and FastAPI applications"),
            ("Meta-analysis aggregates multiple empirical studies", "Reframe", "The product records moderators for a future controlled comparison; no study-level meta-analysis is implemented"),
            ("Empirical validation demonstrates effectiveness", "Remove", "Only functional software verification has been completed"),
            ("Framework outperforms conventional methods", "Remove", "No adequately powered comparative result exists"),
            ("Confusion matrix proves chatbot accuracy", "Reframe", "Matrix summarizes stored task correctness and score threshold agreement; it is not a validated model benchmark"),
            ("Large-scale dataset and generalizable findings", "Remove", "Demo data is simulated and labeled; observed sample size must be reported"),
        ],
        [3600, 1200, 4560],
    )

    add_heading(doc, "References", 1)
    references = [
        "FastAPI Documentation. https://fastapi.tiangolo.com/",
        "SQLAlchemy Documentation. https://docs.sqlalchemy.org/",
        "Alembic Documentation. https://alembic.sqlalchemy.org/",
        "Groq Documentation. https://console.groq.com/docs",
        "Vercel Documentation: Postgres on Vercel. https://vercel.com/docs/postgres",
        "SuperMemo. SM-2 spaced repetition algorithm description.",
        "Council of Europe. Common European Framework of Reference for Languages (CEFR).",
    ]
    add_numbered(doc, references)

    doc.core_properties.title = "LinguaLeap AI - Implementation-Aligned Project Report"
    doc.core_properties.subject = "Truthful technical and evaluation report for LinguaLeap AI"
    doc.core_properties.author = "LinguaLeap AI Project Team"
    doc.core_properties.keywords = "language learning, React, FastAPI, Groq, spaced repetition, learning analytics"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
